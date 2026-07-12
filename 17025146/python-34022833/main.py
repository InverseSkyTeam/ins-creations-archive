#导库
import tkinter as tk
import tkinter.filedialog
import tkinter.messagebox as msg
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import csv
from docx2pdf import convert
import subprocess
import sys

try:
    from docx import Document
    from docx.parts.image import ImagePart
except:
    msg.showinfo("提示","检测到您未安装docx库,即将为您安装,请稍等片刻")
    subprocess.check_output([sys.executable, "-m", "pip", "install", "python-docx","-i","https://pypi.tuna.tsinghua.edu.cn/simple"])
    from docx import Document
    from docx.parts.image import ImagePart

    
#声明全局变量
global docfilename,targetfilename,e,e2,input1,input2
docfilename = None
targetfilename = None


#提取图片函数
def get_all_image(doc):
    document = Document(doc)
    images = []
    for embed, related_part in document.part.related_parts.items():
        if isinstance(related_part, ImagePart):
            images.append(related_part.image)
    return images

def extract_img(target,doc):
    num = 0
    for i in get_all_image(doc):
        with open(target + str(num) + "." + str(i.ext),"wb") as f:
            f.write(i.blob)
        num += 1

def extractimg():
    global input1,input2
    extract_img(input2.get(),input1.get())
    msg.showinfo("完成","提取完成!")


#提取文字函数
def extract_text(target,doc):
    returntext = ""
    wordfile = Document(doc)
    for paragraph in wordfile.paragraphs: 
        returntext += paragraph.text + "\n"
    with open(target + "target.txt","w",encoding="utf-8") as file:
        file.write(returntext)

def extracttext():
    global input1,input2
    extract_text(input2.get(),input1.get())
    msg.showinfo("完成","提取完成!")


#提取表格函数
def extract_form(target,doc):
    doc = Document(doc)
    num = 0
    for table in doc.tables:
        csvtext = []
        for row in table.rows:
            rowlist = []
            for cell in row.cells:
                rowlist.append(cell.text)
            csvtext.append(rowlist)
        with open(target + str() +".csv", mode="w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(csvtext)
        num += 1

def extractform():
    global input1,input2
    extract_form(input2.get(),input1.get())
    msg.showinfo("完成","提取完成!")

#转换为PDF函数
def convert_to_pdf(target,doc):
    convert(doc,target + "target.pdf")

def converttopdf():
    global input1,input2
    convert_to_pdf(input2.get(),input1.get())
    msg.showinfo("完成","提取完成!")

#选择文档及目标路径的函数
def choose_doc():
    global docfilename,e
    docfilename =tkinter.filedialog.askopenfilename(filetypes=[("word文档",".docx")])
    e.set(docfilename)

def choose_target():
    global targetfilename,e2
    targetfilename =tkinter.filedialog.askdirectory() + "/"
    e2.set(targetfilename)


root = ttk.Window(themename="darkly")
root.title("Word文档图片批量提取器")
root.geometry("550x200")

label1 = ttk.Label(root,text="Word文档路径")
label1.place(x = 10,y = 20)
e = ttk.StringVar()
input1 = ttk.Entry(root,width=35,textvariable=e)
input1.place(x = 130,y = 20)
button1 = ttk.Button(root,width = 5,text = "浏览",command=choose_doc)
button1.place(x= 465,y = 20)

label2 = ttk.Label(root,text="目标文件夹路径")
label2.place(x = 10,y = 80)
e2 = ttk.StringVar()
input2 = ttk.Entry(root,width=35,textvariable=e2)
input2.place(x = 130,y = 80)
button2 = ttk.Button(root,width = 5,text = "浏览",command=choose_target)
button2.place(x= 465,y = 80)

button4img = ttk.Button(root,text = "提取图片",width = 7,command=extractimg)
button4img.place(x = 100,y = 150)

button4text = ttk.Button(root,text = "提取文字",width = 7,command=extracttext)
button4text.place(x = 200,y = 150)

button4form = ttk.Button(root,text = "提取表格",width = 7,command=extractform)
button4form.place(x = 300,y = 150)

button4pdf = ttk.Button(root,text = "转换为PDF",width = 8,command=converttopdf)
button4pdf.place(x = 400,y = 150)

root.mainloop()